"""Service for processing documents and extracting text."""

import os
import time
import uuid
from typing import List, Dict, Any, Optional, Tuple
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document

from app.services.embedding_service import generate_embeddings
from app.services.vector_store import store_document_chunks
from app.services.ocr_service import process_image_ocr
from app.models.document import DocumentChunk, DocumentMetadata


def process_document(doc_id: str, file_path: str, title: str, file_type: str, language: str = "english") -> Dict[str, Any]:
    """Process a document file and store its chunks in the vector database.
    
    Args:
        doc_id: Unique identifier for the document
        file_path: Path to the document file
        title: Document title
        file_type: File extension (e.g., ".pdf", ".docx")
        language: Document language
        
    Returns:
        Dictionary with processing results
    """
    start_time = time.time()
    
    try:
        # Extract text based on file type
        if file_type.lower() == ".pdf":
            text, metadata = extract_text_from_pdf(file_path)
        elif file_type.lower() == ".docx":
            text, metadata = extract_text_from_docx(file_path)
        elif file_type.lower() == ".txt":
            text, metadata = extract_text_from_text(file_path)
        elif file_type.lower() in [".jpg", ".jpeg", ".png"]:
            text, metadata = extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Update metadata
        metadata["doc_id"] = doc_id
        metadata["title"] = title
        
        # For image files, use the detected language from OCR
        if file_type.lower() in [".jpg", ".jpeg", ".png"] and "detected_language" in metadata:
            metadata["language"] = metadata["detected_language"]
        else:
            metadata["language"] = language
        
        # Calculate word count
        word_count = len(text.split())
        metadata["word_count"] = word_count
        
        # Split text into chunks
        chunks = split_text_into_chunks(text, chunk_size=1000, chunk_overlap=200)
        
        # Create document chunks
        doc_chunks = []
        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{doc_id}-chunk-{i}"
            chunk_metadata = {
                "chunk_id": chunk_id,
                "doc_id": doc_id,
                "chunk_index": i,
                "source": f"Chunk {i+1} of {len(chunks)}",
            }
            
            doc_chunks.append(
                DocumentChunk(
                    id=chunk_id,
                    doc_id=doc_id,
                    text=chunk_text,
                    metadata=chunk_metadata,
                )
            )
        
        # Generate embeddings for chunks
        doc_chunks_with_embeddings = generate_embeddings(doc_chunks)
        
        # Store chunks in vector database
        store_document_chunks(doc_chunks_with_embeddings)
        
        # Create document metadata
        doc_metadata = DocumentMetadata(
            id=doc_id,
            title=title,
            upload_date=time.strftime("%Y-%m-%dT%H:%M:%S"),
            file_type=file_type.replace(".", ""),
            page_count=metadata.get("page_count"),
            chunk_count=len(chunks),
            language=language,
        )
        
        # Calculate processing time
        processing_time = time.time() - start_time
        
        return {
            "status": "success",
            "doc_id": doc_id,
            "title": title,
            "chunk_count": len(chunks),
            "word_count": word_count,
            "page_count": metadata.get("page_count", 1),
            "processing_time": processing_time,
            "metadata": doc_metadata.dict(),
        }
    
    except Exception as e:
        # Log the error
        print(f"Error processing document {doc_id}: {str(e)}")
        
        # Return error information
        return {
            "status": "error",
            "doc_id": doc_id,
            "title": title,
            "error": str(e),
        }


def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Tuple of (extracted text, metadata dictionary)
    """
    with open(file_path, "rb") as file:
        pdf_reader = PyPDF2.PdfReader(file)
        num_pages = len(pdf_reader.pages)
        
        # Extract text from each page
        text = ""
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n\n"
        
        # Extract metadata
        metadata = {
            "page_count": num_pages,
            "file_size": os.path.getsize(file_path),
        }
        
        # Try to extract PDF metadata if available
        if pdf_reader.metadata:
            for key, value in pdf_reader.metadata.items():
                if value and key.startswith("/"):
                    metadata[key[1:].lower()] = value
    
    return text, metadata


def extract_text_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple of (extracted text, metadata dictionary)
    """
    doc = Document(file_path)
    
    # Extract text from paragraphs
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Extract metadata
    metadata = {
        "file_size": os.path.getsize(file_path),
        "paragraph_count": len(doc.paragraphs),
    }
    
    # Try to extract document properties if available
    try:
        core_properties = doc.core_properties
        metadata["author"] = core_properties.author
        metadata["created"] = core_properties.created
        metadata["modified"] = core_properties.modified
        metadata["title"] = core_properties.title
        metadata["subject"] = core_properties.subject
    except Exception:
        pass
    
    return text, metadata


def extract_text_from_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a plain text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        Tuple of (extracted text, metadata dictionary)
    """
    with open(file_path, "r", encoding="utf-8") as file:
        text = file.read()
    
    # Extract metadata
    metadata = {
        "file_size": os.path.getsize(file_path),
        "line_count": text.count("\n") + 1,
    }
    
    return text, metadata


def extract_text_from_image(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from an image file using OCR.
    
    Args:
        file_path: Path to the image file
        
    Returns:
        Tuple of (extracted text, metadata)
    """
    try:
        # Generate a unique OCR ID
        ocr_id = str(uuid.uuid4())
        
        # Process the image with OCR (language will be auto-detected)
        ocr_result = process_image_ocr(
            ocr_id=ocr_id,
            file_path=file_path,
            language=None,  # Auto-detect language
            enhance_image=True,
            use_easyocr=True
        )
        
        # Extract text and metadata from OCR result
        if ocr_result["status"] == "success":
            text = ocr_result["text"]
            
            # Create metadata
            metadata = {
                "page_count": 1,  # Image files are considered as a single page
                "word_count": len(text.split()),
                "ocr_confidence": ocr_result.get("confidence", 0),
                "detected_language": ocr_result.get("language", "english"),
                "ocr_engine": ocr_result.get("engine", "EasyOCR"),
                "processing_time": ocr_result.get("processing_time", 0),
            }
            
            return text, metadata
        else:
            error_message = ocr_result.get("error", "Unknown OCR error")
            print(f"Error in OCR processing: {error_message}")
            return "", {"error": error_message}
    
    except Exception as e:
        print(f"Error extracting text from image file: {e}")
        return "", {"error": str(e)}


def extract_text_from_url(doc_id: str, url: str, title: str) -> Dict[str, Any]:
    """Extract text from a web page URL.
    
    Args:
        doc_id: Unique identifier for the document
        url: Web page URL
        title: Document title
        
    Returns:
        Dictionary with processing results
    """
    try:
        # Fetch the web page
        response = requests.get(url, headers={"User-Agent": "SmartDocQ/1.0"})
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.text, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text(separator="\n")
        
        # Clean up text (remove excessive whitespace)
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        
        # Extract metadata
        metadata = {
            "source_url": url,
            "title": title or soup.title.string if soup.title else url,
            "doc_id": doc_id,
        }
        
        # Process the extracted text as a document
        return process_web_text(doc_id, text, metadata)
    
    except Exception as e:
        # Log the error
        print(f"Error processing URL {url}: {str(e)}")
        
        # Return error information
        return {
            "status": "error",
            "doc_id": doc_id,
            "title": title or url,
            "error": str(e),
        }


def process_web_text(doc_id: str, text: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Process extracted web page text.
    
    Args:
        doc_id: Unique identifier for the document
        text: Extracted text from web page
        metadata: Metadata dictionary
        
    Returns:
        Dictionary with processing results
    """
    start_time = time.time()
    
    try:
        # Split text into chunks
        chunks = split_text_into_chunks(text, chunk_size=1000, chunk_overlap=200)
        
        # Create document chunks
        doc_chunks = []
        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{doc_id}-chunk-{i}"
            chunk_metadata = {
                "chunk_id": chunk_id,
                "doc_id": doc_id,
                "chunk_index": i,
                "source": f"Web: {metadata.get('source_url')}, Chunk {i+1} of {len(chunks)}",
            }
            
            doc_chunks.append(
                DocumentChunk(
                    id=chunk_id,
                    doc_id=doc_id,
                    text=chunk_text,
                    metadata=chunk_metadata,
                )
            )
        
        # Generate embeddings for chunks
        doc_chunks_with_embeddings = generate_embeddings(doc_chunks)
        
        # Store chunks in vector database
        store_document_chunks(doc_chunks_with_embeddings)
        
        # Create document metadata
        doc_metadata = DocumentMetadata(
            id=doc_id,
            title=metadata.get("title"),
            upload_date=time.strftime("%Y-%m-%dT%H:%M:%S"),
            file_type="web",
            chunk_count=len(chunks),
            source_url=metadata.get("source_url"),
        )
        
        # Calculate processing time
        processing_time = time.time() - start_time
        
        return {
            "status": "success",
            "doc_id": doc_id,
            "title": metadata.get("title"),
            "chunk_count": len(chunks),
            "processing_time": processing_time,
            "metadata": doc_metadata.dict(),
        }
    
    except Exception as e:
        # Log the error
        print(f"Error processing web text for {doc_id}: {str(e)}")
        
        # Return error information
        return {
            "status": "error",
            "doc_id": doc_id,
            "title": metadata.get("title"),
            "error": str(e),
        }


def split_text_into_chunks(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks.
    
    Args:
        text: Text to split
        chunk_size: Maximum size of each chunk in characters
        chunk_overlap: Overlap between chunks in characters
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
    
    # Split text into sentences (simple approach)
    sentences = [s.strip() for s in text.replace("\n", " ").split(".") if s.strip()]
    
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence_size = len(sentence)
        
        # If adding this sentence would exceed chunk size, finalize the current chunk
        if current_size + sentence_size > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk) + ".")
            
            # Keep some sentences for overlap
            overlap_size = 0
            overlap_sentences = []
            
            # Start from the end and work backwards to create overlap
            for s in reversed(current_chunk):
                if overlap_size + len(s) <= chunk_overlap:
                    overlap_sentences.insert(0, s)
                    overlap_size += len(s)
                else:
                    break
            
            current_chunk = overlap_sentences
            current_size = overlap_size
        
        # Add the current sentence to the chunk
        current_chunk.append(sentence)
        current_size += sentence_size
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(" ".join(current_chunk) + ".")
    
    return chunks