"""Utility functions for chat functionality in the frontend."""

import streamlit as st
import requests
import os
import time


def display_chat_history(chat_history):
    """Display the chat history with proper formatting.
    
    Args:
        chat_history (list): List of chat message dictionaries
    """
    if not chat_history:
        st.info("No messages yet. Start by asking a question about your document.")
        return
    
    # Display each message in the chat history
    for message in chat_history:
        display_chat_message(message)


def display_chat_message(message):
    """Display a single chat message with proper formatting.
    
    Args:
        message (dict): Chat message dictionary
    """
    # Determine message class based on role
    message_class = "user-message" if message["role"] == "user" else "assistant-message"
    
    # Create a container for the message
    with st.container():
        # Apply custom CSS class with stronger text color and preserved whitespace
        # Force dark text inside light chat bubbles for readability
        st.markdown(
            f"<div class='chat-message {message_class}' style='color:#0b1220; white-space: pre-wrap;'>"
            f"<div style='opacity:.9; font-weight:600; margin-bottom:.25rem;'>{'You' if message['role'] == 'user' else 'AI'}:</div>"
            f"<div style='line-height:1.5;'>{message['content']}</div>"
            f"</div>",
            unsafe_allow_html=True,
        )
        
        # Display sources if available
        if message.get("sources"):
            with st.expander("Sources"):
                for i, source in enumerate(message["sources"]):
                    st.markdown(f"**Source {i+1}:** {source['text']}")
                    st.markdown(f"<span class='source-info'>{source['source']}</span>", unsafe_allow_html=True)
        
        # Display timestamp if available
        if message.get("timestamp"):
            st.markdown(f"<span class='source-info'>{format_timestamp(message['timestamp'])}</span>", 
                       unsafe_allow_html=True)
        
        # Display feedback options for assistant messages
        if message["role"] == "assistant" and not message.get("feedback_submitted"):
            with st.container():
                # Use timestamp and index in chat history to ensure unique keys
                message_idx = next((i for i, m in enumerate(st.session_state.chat_history) if m is message), 0)
                unique_id = f"{message_idx}_{message.get('timestamp', time.time())}"
                
                col1, col2, col3 = st.columns([1, 1, 3])
                with col1:
                    if st.button("👍", key=f"thumbs_up_{unique_id}"):
                        submit_feedback(message, "positive")
                with col2:
                    if st.button("👎", key=f"thumbs_down_{unique_id}"):
                        submit_feedback(message, "negative")
                with col3:
                    if st.button("Report Issue", key=f"report_{unique_id}"):
                        show_feedback_form(message)


def add_message_to_history(role, content, sources=None):
    """Add a new message to the chat history.
    
    Args:
        role (str): Message role ("user" or "assistant")
        content (str): Message content
        sources (list, optional): List of source dictionaries
    """
    # Create message dictionary
    message = {
        "role": role,
        "content": content,
        "timestamp": time.time(),
    }
    
    # Add sources if provided
    if sources:
        message["sources"] = sources
    
    # Add to session state
    st.session_state.chat_history.append(message)


def submit_feedback(message, feedback_type):
    """Submit feedback for a message.
    
    Args:
        message (dict): Message dictionary
        feedback_type (str): Type of feedback ("positive" or "negative")
    """
    # Submit feedback to backend API
    try:
        API_URL = os.getenv("API_URL", "http://localhost:8000/api")
        rating = 5 if feedback_type == "positive" else 2
        payload = {
            "question_id": str(message.get("timestamp", "")),
            "rating": rating,
            "comment": None,
            "is_error_report": False,
        }
        requests.post(f"{API_URL}/feedback/submit", json=payload, timeout=6)
    except Exception:
        pass

    # Mark locally
    message["feedback_submitted"] = True
    message["feedback_type"] = feedback_type

    # Attempt lightweight refinement on negative feedback
    if feedback_type == "negative":
        try:
            # Find last user question preceding this assistant message
            idx = next((i for i, m in enumerate(st.session_state.chat_history) if m is message), -1)
            last_user = None
            for m in reversed(st.session_state.chat_history[:idx]):
                if m.get("role") == "user":
                    last_user = m.get("content")
                    break
            doc = st.session_state.get("current_document")
            if last_user and doc:
                API_URL = os.getenv("API_URL", "http://localhost:8000/api")
                data = {
                    "question": f"Improve and correct the previous answer based on user feedback for: {last_user}",
                    "document_id": doc["id"],
                }
                resp = requests.post(f"{API_URL}/chat/question", json=data, timeout=15)
                if resp.status_code == 200:
                    new_answer = resp.json().get("answer") or message.get("content")
                    message["content"] = new_answer
        except Exception:
            pass

    st.success("Thank you for your feedback!")
    st.rerun()


def show_feedback_form(message):
    """Show a form to collect detailed feedback.
    
    Args:
        message (dict): Message dictionary
    """
    # Get unique ID for the message
    message_idx = next((i for i, m in enumerate(st.session_state.chat_history) if m is message), 0)
    unique_id = f"{message_idx}_{message.get('timestamp', time.time())}"
    
    # Create a form for feedback
    with st.form(key=f"feedback_form_{unique_id}"):
        st.write("**Report an Issue**")
        issue_type = st.selectbox(
            "Issue Type",
            ["Incorrect Information", "Missing Information", "Irrelevant Answer", "Other"],
        )
        details = st.text_area("Details", height=100)
        submit_button = st.form_submit_button("Submit Report")
    
    if submit_button:
        # Try backend API for issue reporting
        try:
            API_URL = os.getenv("API_URL", "http://localhost:8000/api")
            requests.post(
                f"{API_URL}/feedback/report-error",
                json=None,
                data=None,
                timeout=6,
            )
        except Exception:
            # If backend is unavailable and report cannot be sent, hide the feature next render
            st.session_state["disable_issue_report"] = True
        # Mark locally
        message["feedback_submitted"] = True
        message["feedback_type"] = "issue"
        message["feedback_details"] = {"issue_type": issue_type, "details": details}
        st.success("Thank you for your feedback! We'll review this issue.")
        st.rerun()


def format_timestamp(timestamp):
    """Format a timestamp for display.
    
    Args:
        timestamp (float): Unix timestamp
        
    Returns:
        str: Formatted timestamp string
    """
    return time.strftime("%I:%M %p", time.localtime(timestamp))


def clear_chat_history():
    """Clear the chat history."""
    st.session_state.chat_history = []


def export_chat_history(format_type="pdf"):
    """Export the chat history to a PDF or DOCX file and return bytes.
    
    Args:
        format_type (str): "pdf" or "docx"
    Returns:
        bytes: Binary file content
    """
    messages = st.session_state.get("chat_history", [])
    transcript_lines = []
    for m in messages:
        speaker = "You" if m.get("role") == "user" else "AI"
        time_str = format_timestamp(m.get("timestamp", 0))
        transcript_lines.append(f"{speaker} ({time_str})\n{m.get('content','')}\n")
    transcript = "\n\n".join(transcript_lines) or "(No messages)"

    if format_type.lower() == "docx":
        from docx import Document
        from io import BytesIO
        doc = Document()
        doc.add_heading("SmartDocQ Chat Export", level=1)
        for block in transcript.split("\n\n"):
            p = doc.add_paragraph(block)
            p.paragraph_format.space_after = 12
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Default PDF generation
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from io import BytesIO
    styles = getSampleStyleSheet()
    normal = styles["BodyText"]
    normal.fontName = "Helvetica"
    normal.fontSize = 10
    story = []
    title = Paragraph("<b>SmartDocQ Chat Export</b>", styles["Title"])
    story.append(title)
    story.append(Spacer(1, 6*mm))
    for block in transcript.split("\n\n"):
        story.append(Paragraph(block.replace("\n", "<br/>"), normal))
        story.append(Spacer(1, 4*mm))
    buf = BytesIO()
    SimpleDocTemplate(buf, pagesize=A4).build(story)
    return buf.getvalue()